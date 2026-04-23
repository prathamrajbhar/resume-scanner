import os
import io
import re
import json
import logging
import string
import fitz  # PyMuPDF
import ollama
from typing import List, Dict, Set, Tuple, Optional
from pathlib import Path
from app.core.config import settings
import numpy as np

# Optional heavy imports like SentenceTransformer are moved to methods for lazy loading

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class NLPService:
    _instance = None
    _bert_model = None

    def __new__(cls, *args, **kwargs):
        if not cls._instance:
            cls._instance = super(NLPService, cls).__new__(cls)
        return cls._instance

    def __init__(self):
        # Prevent re-initialization
        if hasattr(self, "_initialized") and self._initialized:
            return
            
        # Initialize Ollama settings
        self.model_name = settings.OLLAMA_MODEL
        self.base_url = settings.OLLAMA_BASE_URL
        self.client = ollama.AsyncClient(host=self.base_url)
        logger.info(f"Initialized singleton NLPService with Ollama AsyncClient: {self.model_name} at {self.base_url}")
        self._initialized = True

    @classmethod
    def get_bert_model(cls):
        if cls._bert_model is None:
            logger.info("Loading BERT model...")
            from sentence_transformers import SentenceTransformer
            cls._bert_model = SentenceTransformer('all-MiniLM-L6-v2')
        return cls._bert_model

    def extract_text_from_bytes(self, content: bytes, filename: str) -> str:
        """Extracts raw text from various file formats."""
        ext = filename.split('.')[-1].lower()
        text = ""
        
        try:
            if ext == 'pdf':
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text += page.get_text()
                doc.close()
            elif ext in ['doc', 'docx']:
                try:
                    import docx
                    import io
                    doc = docx.Document(io.BytesIO(content))
                    
                    # Extract from paragraphs
                    paragraphs = [p.text for p in doc.paragraphs]
                    
                    # Extract from tables (very common in resumes)
                    tables = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                tables.append(" | ".join(row_text))
                    
                    text = "\n".join(paragraphs + tables)
                    logger.info(f"Extracted {len(text)} characters from DOCX (including tables). Preview: {text[:100]}...")
                except Exception as e:
                    logger.error(f"Error parsing DOCX: {e}")
                    text = ""
            elif ext in ['csv']:
                import pandas as pd
                df = pd.read_csv(io.BytesIO(content))
                text = df.to_string()
            elif ext in ['xlsx', 'xls']:
                import pandas as pd
                df = pd.read_excel(io.BytesIO(content))
                text = df.to_string()
            else:
                text = content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            
        return text.strip()

    async def _call_ollama_json(self, prompt: str, max_retries: int = 3) -> Dict:
        """Helper to call Ollama and parse JSON response with retry logic."""
        import asyncio
        
        for attempt in range(max_retries):
            try:
                response = await self.client.chat(
                    model=self.model_name,
                    messages=[{'role': 'user', 'content': prompt}],
                    format='json',
                    options={'temperature': 0}
                )
                
                content = response.get('message', {}).get('content', '')
                if not content:
                    logger.warning("Ollama returned an empty response content.")
                    return {}
                
                return json.loads(content)
            except Exception as e:
                logger.error(f"Ollama API error (Attempt {attempt+1}/{max_retries}): {str(e)}")
                if attempt == max_retries - 1:
                    return {}
                await asyncio.sleep(2 ** attempt)
        return {}

    async def extract_candidate_profile(self, text: str, filename: str = "") -> Dict:
        """Exclusively Ollama-based profile extraction."""
        # Clean text: remove excessive whitespace and limit to a reasonable chunk
        # 30,000 characters is plenty for even long CVs.
        clean_text = re.sub(r'\s+', ' ', text).strip()
        
        prompt = f"""
        Extract professional information from the following resume text.
        Return the result as a JSON object with these fields:
        - full_name: string
        - email: string
        - phone: string
        - location: string
        - skills: list of strings (technical skills)
        - soft_skills: list of strings
        - degrees: list of objects (degree, field, college, cgpa, location)
        - experience_list: list of objects (role, company, duration, description)
        - projects: list of strings
        - certifications: list of strings
        - awards: list of strings
        - total_experience_years: float (number of years)
        - projects_count: integer
        - education: string (highest degree summary)
        - university: string (primary university)
        - cgpa: float (normalized to 10.0 scale)
        - sgpa: float
        - internships: list of strings

        Resume Text:
        {clean_text[:30000]}
        """
        
        extracted = await self._call_ollama_json(prompt)
        
        # Ensure all fields exist to avoid frontend crashes
        defaults = {
            'full_name': None, 'email': None, 'phone': None, 'location': None,
            'skills': [], 'soft_skills': [], 'degrees': [], 'experience_list': [],
            'projects': [], 'certifications': [], 'awards': [],
            'total_experience_years': 0.0, 'projects_count': 0,
            'education': None, 'university': None, 'cgpa': None, 'sgpa': None,
            'internships': []
        }
        
        # Merge with defaults
        profile = {**defaults, **extracted}
        
        # Compatibility aliases for existing frontend/backend logic
        profile['normalized_skills'] = profile['skills']
        profile['total_experience'] = profile['total_experience_years']
        profile['cgpa_or_percentage'] = profile['cgpa']
        
        return profile

    async def extract_skills(self, text: str) -> Set[str]:
        """Extract skills using Ollama."""
        prompt = f"List all technical and soft skills mentioned in this text as a JSON array of strings: {text[:10000]}"
        response = await self._call_ollama_json(prompt)
        return set(response) if isinstance(response, list) else set()

    async def analyze_candidate(self, resume_text: str, job_description: str, model_type: str = "ensemble") -> Dict:
        """Intelligent Ollama-based resume-to-JD matching with full factor extraction."""
        prompt = f"""
        Compare the following resume against the job description.
        Evaluate the candidate on technical skills, experience relevance, education, and soft skills.
        
        Return a JSON object with:
        - score: float (0-100 overall fit score)
        - matching_skills: list of strings
        - missing_skills: list of strings
        - explanation: string (brief reason for the score)
        
        - academic_data: {{
            "degree": string,
            "university": string,
            "cgpa": float,
            "percentage": float,
            "academic_score": float (0-100)
          }}
        - experience_data: {{
            "total_years": float,
            "relevant_years": float,
            "projects_count": integer,
            "experience_score": float (0-100)
          }}
        - soft_skills_data: {{
            "communication": float (0-100),
            "leadership": float (0-100),
            "teamwork": float (0-100),
            "problem_solving": float (0-100),
            "soft_skill_score": float (0-100)
          }}
        - additional_data: {{
            "skill_match_score": float (0-100),
            "project_score": float (0-100)
          }}

        Job Description:
        {job_description}

        Resume:
        {resume_text[:20000]}
        """
        
        analysis = await self._call_ollama_json(prompt)
        
        # Default structure if LLM fails
        if not analysis:
            return {
                "score": 0.0,
                "matching_skills": [],
                "missing_skills": [],
                "explanation": "Analysis failed",
                "academic_data": {"academic_score": 0},
                "experience_data": {"experience_score": 0},
                "soft_skills_data": {"soft_skill_score": 0},
                "additional_data": {"skill_match_score": 0, "project_score": 0}
            }
            
        # Enrich with exhaustive resume skills for completeness
        if "skills" not in analysis:
            analysis["skills"] = list(await self.extract_skills(resume_text))
        
        return analysis

    def score_tfidf(self, resume_text: str, job_description: str) -> float:
        """Legacy support for TF-IDF scoring."""
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        if not resume_text or not job_description: return 0.0
        vectorizer = TfidfVectorizer(stop_words='english')
        tfidf = vectorizer.fit_transform([resume_text, job_description])
        return float(cosine_similarity(tfidf[0:1], tfidf[1:2])[0][0])

    def score_bert(self, resume_text: str, job_description: str) -> float:
        """Legacy support for BERT scoring."""
        from sklearn.metrics.pairwise import cosine_similarity
        model = self.get_bert_model()
        embeddings = model.encode([resume_text, job_description])
        return float(cosine_similarity([embeddings[0]], [embeddings[1]])[0][0])

    async def extract_location(self, text: str) -> Optional[str]:
        """Extract location using Ollama."""
        prompt = f"Extract the current location of the person from this text. Return as JSON: {{\"location\": \"string\"}}. Text: {text[:5000]}"
        res = await self._call_ollama_json(prompt)
        return res.get("location")

    async def extract_candidate_name(self, text: str, filename: str = "") -> Optional[str]:
        """Extract candidate name using Ollama."""
        prompt = f"Extract the full name of the candidate from this resume text. Return as JSON: {{\"name\": \"string\"}}. Text: {text[:5000]}"
        res = await self._call_ollama_json(prompt)
        return res.get("name")

    async def get_professional_insights(self, combined_text: str, role_title: Optional[str] = None) -> Dict:
        """Generates deep professional insights using Ollama."""
        prompt = f"""
        Analyze the following professional profile (Resume + LinkedIn).
        Job Role Context: {role_title or "General"}
        
        Evaluate the candidate on 4 dimensions (0-100 score):
        1. Communication: Clarity, structure, and professional tone.
        2. Domain Fit: Alignment of skills and experience with the target role.
        3. Learning Ability: Evidence of certifications, continuous learning, or rapid growth.
        4. Career Stability: Tenure patterns, role transitions, and progression.

        For each dimension, provide:
        - score: float
        - explanation: string
        - reasons: list of strings (bullet points)
        - evidence: list of strings (snippets from text)

        Return as a structured JSON object.

        Profile Text:
        {combined_text[:20000]}
        """
        return await self._call_ollama_json(prompt)
